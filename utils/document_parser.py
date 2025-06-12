import io
import os
import tempfile
from PyPDF2 import PdfReader
from docx import Document
import pdfplumber
try:
    import pytesseract
    from PIL import Image
    import fitz  # PyMuPDF
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("Warning: OCR dependencies not available. Install pytesseract, Pillow, and PyMuPDF for OCR support.")

class DocumentParser:
    """Parse different document formats to extract text content"""

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']

    def extract_text_from_file(self, file_obj, filename):
        """Extract text content from uploaded file object"""
        try:
            file_extension = os.path.splitext(filename)[1].lower()

            if file_extension == '.pdf':
                result = self._extract_from_pdf(file_obj)
            elif file_extension in ['.docx', '.doc']:
                result = self._extract_from_docx(file_obj)
            elif file_extension == '.txt':
                result = self._extract_from_txt(file_obj)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")

            # Ensure we always return a string
            if isinstance(result, bytes):
                return result.decode('utf-8', errors='ignore')
            return str(result)

        except Exception as e:
            print(f"Error parsing {filename}: {str(e)}")
            # Try fallback text extraction
            try:
                file_obj.seek(0)
                content = file_obj.read()
                if isinstance(content, bytes):
                    return content.decode('utf-8', errors='ignore')
                return str(content)
            except:
                return f"Error: Could not parse {filename}. Please ensure the file is not corrupted."

    def _extract_from_pdf(self, file_obj):
        """Extract text from PDF file with OCR fallback"""
        text_content = []

        try:
            # Try with pdfplumber first (better for complex PDFs)
            file_obj.seek(0)
            with pdfplumber.open(file_obj) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_content.append(page_text)

            # If we got substantial text, return it
            combined_text = "\n".join(text_content)
            if len(combined_text.strip()) > 50:  # Minimum text threshold
                return combined_text
        except Exception as e:
            print(f"pdfplumber failed: {e}")

        try:
            # Fallback to PyPDF2
            file_obj.seek(0)
            pdf_reader = PdfReader(file_obj)
            text_content = []

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_content.append(page_text)

            combined_text = "\n".join(text_content)
            if len(combined_text.strip()) > 50:
                return combined_text
        except Exception as e:
            print(f"PyPDF2 failed: {e}")

        # If text extraction failed or returned minimal text, try OCR
        if OCR_AVAILABLE:
            try:
                return self._extract_with_ocr(file_obj)
            except Exception as e:
                print(f"OCR failed: {e}")

        return "Unable to extract text from PDF. The file may be image-based or corrupted."

    def _extract_with_ocr(self, file_obj):
        """Extract text using OCR for image-based PDFs"""
        if not OCR_AVAILABLE:
            return "OCR not available"

        text_content = []
        file_obj.seek(0)

        # Save to temporary file for PyMuPDF
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_file.write(file_obj.read())
            temp_file_path = temp_file.name

        try:
            # Open PDF with PyMuPDF
            pdf_document = fitz.open(temp_file_path)

            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)

                # Convert page to image
                pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))  # Higher resolution
                img_data = pix.tobytes("png")

                # Convert to PIL Image
                image = Image.open(io.BytesIO(img_data))

                # Perform OCR
                page_text = pytesseract.image_to_string(image, config='--psm 6')
                if page_text.strip():
                    text_content.append(page_text)

            pdf_document.close()
            return "\n".join(text_content) if text_content else "No text found via OCR"

        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_file_path)
            except:
                pass

    def _extract_from_docx(self, file_obj):
        """Extract text from DOCX file"""
        try:
            # Save to temporary file since python-docx needs file path
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                file_obj.seek(0)
                temp_file.write(file_obj.read())
                temp_file_path = temp_file.name

            try:
                doc = Document(temp_file_path)
                text_content = []

                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)

                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_content.append(cell.text)

                return "\n".join(text_content) if text_content else "No text found in document"

            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except:
                    pass

        except Exception as e:
            return f"Error reading DOCX file: {str(e)}"

    def _extract_from_txt(self, file_obj):
        """Extract text from TXT file"""
        try:
            file_obj.seek(0)
            content = file_obj.read()

            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

            for encoding in encodings:
                try:
                    if isinstance(content, bytes):
                        return content.decode(encoding)
                    else:
                        return content
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, use ignore errors
            if isinstance(content, bytes):
                return content.decode('utf-8', errors='ignore')
            else:
                return content

        except Exception as e:
            return f"Error reading text file: {str(e)}"

    def is_supported_format(self, filename):
        """Check if file format is supported"""
        file_extension = os.path.splitext(filename)[1].lower()
        return file_extension in self.supported_formats

    def get_file_info(self, filename):
        """Get information about the file"""
        file_extension = os.path.splitext(filename)[1].lower()

        format_info = {
            '.pdf': 'PDF Document',
            '.docx': 'Microsoft Word Document',
            '.doc': 'Microsoft Word Document (Legacy)',
            '.txt': 'Plain Text File'
        }

        return {
            'extension': file_extension,
            'format': format_info.get(file_extension, 'Unknown Format'),
            'supported': file_extension in self.supported_formats
        }